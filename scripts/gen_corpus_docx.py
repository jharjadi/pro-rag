"""Generate expanded DOCX documents for Phase 3b.5 corpus (docs 6, 7, 8, 14, 15)."""
from __future__ import annotations
from pathlib import Path
from docx import Document

OUTPUT_DIR = Path(__file__).parent.parent / "data" / "test-corpus"


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value


def generate_leave_benefits() -> None:
    """Doc 6: Annual Leave & Benefits Summary (table-heavy)."""
    doc = Document()
    doc.add_heading("Acme Corp Annual Leave & Benefits Summary", level=0)
    doc.add_paragraph("Document ID: HR-BEN-002 | Version: 5.0 | Effective: January 1, 2026 | Owner: HR")
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "Comprehensive summary of leave entitlements and benefits. Full-time employees "
        "eligible from day one. Part-time (20+ hrs/week) get pro-rated benefits."
    )
    doc.add_heading("2. Leave by Tenure", level=1)
    _add_table(doc, ["Tenure", "PTO", "Sick", "Personal", "Total"], [
        ["0-1 years", "15 days", "10 days", "2 days", "27 days"],
        ["1-3 years", "20 days", "10 days", "3 days", "33 days"],
        ["3-5 years", "25 days", "12 days", "3 days", "40 days"],
        ["5-7 years", "25 days", "12 days", "4 days", "41 days"],
        ["7-10 years", "30 days", "15 days", "4 days", "49 days"],
        ["10+ years", "30 days", "15 days", "5 days", "50 days"],
    ])
    doc.add_heading("3. Health Insurance", level=1)
    doc.add_paragraph("Three plan options. Enroll within 30 days of start or during open enrollment.")
    _add_table(doc, ["Feature", "PPO", "HMO", "HDHP+HSA"], [
        ["Premium (Employee)", "$250/mo", "$180/mo", "$120/mo"],
        ["Premium (Family)", "$750/mo", "$540/mo", "$360/mo"],
        ["Deductible", "$500", "$250", "$2,800"],
        ["Out-of-Pocket Max", "$4,000", "$3,000", "$7,050"],
        ["Primary Care", "$25 copay", "$15 copay", "20% after ded."],
        ["Specialist", "$50 copay", "$30 copay", "20% after ded."],
        ["ER", "$250 copay", "$200 copay", "20% after ded."],
        ["Rx Generic", "$10", "$5", "20% after ded."],
        ["HSA Contribution", "N/A", "N/A", "$1,500/yr employer"],
    ])
    doc.add_heading("4. 401(k)", level=1)
    _add_table(doc, ["Feature", "Details"], [
        ["Limit", "$23,000/year (2026)"],
        ["Catch-Up (50+)", "+$7,500/year"],
        ["Match", "100% first 4% + 50% next 2% = up to 5%"],
        ["Vesting", "25%/year, full at 4 years"],
        ["Options", "15 index, 5 target-date, self-directed"],
    ])
    doc.add_heading("5. Insurance", level=1)
    _add_table(doc, ["Type", "Benefit", "Cost"], [
        ["Basic Life", "2x salary (max $500K)", "Free"],
        ["STD", "60% salary, 26 weeks", "Free"],
        ["LTD", "60% salary, $15K/mo max", "Free"],
        ["AD&D", "2x salary", "Free"],
    ])
    doc.add_heading("6. Other Benefits", level=1)
    _add_table(doc, ["Benefit", "Details", "Eligibility"], [
        ["Tuition", "$5,250/year", "After 1 year"],
        ["ESPP", "15% discount, $25K/yr", "After 6 months"],
        ["Gym", "$50/month", "Immediate"],
        ["Referral", "$2,500-$5,000/hire", "Immediate"],
    ])
    doc.save(OUTPUT_DIR / "leave_benefits_summary.docx")
    print("  Created: leave_benefits_summary.docx")


def generate_sdlc() -> None:
    """Doc 7: Software Development Lifecycle."""
    doc = Document()
    doc.add_heading("Acme Corp Software Development Lifecycle", level=0)
    doc.add_paragraph("Document ID: ENG-SDLC-001 | Version: 2.3 | Effective: February 1, 2026 | Owner: VP Engineering")
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph("Defines SDLC practices for consistent quality, security, and reliability.")
    doc.add_heading("2. Methodology", level=1)
    doc.add_paragraph(
        "Agile with two-week sprints. Sprint planning, review, retro. Jira for PM, "
        "GitHub for code. All work tracked with story points and acceptance criteria."
    )
    doc.add_heading("3. Code Review", level=1)
    doc.add_paragraph(
        "All changes need peer review. PRs require 2 approvals (1 Staff+ engineer). "
        "24-hour turnaround. Gates: tests, coverage, lint, SAST, CI build."
    )
    doc.add_heading("4. Testing", level=1)
    _add_table(doc, ["Level", "Coverage", "Owner", "Frequency"], [
        ["Unit", "80% lines", "Developer", "Every commit"],
        ["Integration", "Critical paths", "Dev + QA", "Every PR"],
        ["E2E", "Happy + error", "QA", "Nightly"],
        ["Performance", "P95 benchmarks", "Perf Team", "Weekly"],
        ["Security", "OWASP Top 10", "Security", "Weekly"],
    ])
    doc.add_heading("5. Deployment", level=1)
    doc.add_paragraph(
        "CI/CD via GitHub Actions + ArgoCD. Stages: build, test, staging (auto on main), "
        "staging validation, prod (release manager approval), prod validation (30 min). "
        "No manual prod deploys. Auto-rollback if health checks fail within 5 minutes."
    )
    doc.add_heading("6. Environments", level=1)
    _add_table(doc, ["Env", "Purpose", "Data", "Deploy"], [
        ["Dev", "Local testing", "Synthetic", "Continuous"],
        ["CI", "Automated tests", "Fixtures", "Every commit"],
        ["Staging", "Pre-prod", "Anonymized prod", "On main merge"],
        ["Production", "Live", "Real data", "On approval"],
    ])
    doc.add_heading("7. Security", level=1)
    doc.add_paragraph(
        "Shift-left. Annual secure coding training. Dependabot + Snyk. Critical vulns "
        "patched in 7 days. No secrets in code (pre-commit hooks). Vault, 90-day rotation."
    )
    doc.save(OUTPUT_DIR / "software_development_lifecycle.docx")
    print("  Created: software_development_lifecycle.docx")


def generate_data_retention() -> None:
    """Doc 8: Data Retention Policy."""
    doc = Document()
    doc.add_heading("Acme Corp Data Retention Policy", level=0)
    doc.add_paragraph("Document ID: LEG-DRP-001 | Version: 3.1 | Effective: January 1, 2026 | Owner: Legal + CISO")
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "Requirements for retaining and disposing of business records. Ensures legal "
        "compliance, supports operations, reduces storage costs, minimizes risk."
    )
    doc.add_heading("2. Retention Schedule", level=1)
    _add_table(doc, ["Category", "Examples", "Min", "Max", "Basis"], [
        ["Financial", "Ledger, AP/AR, tax", "7 yr", "10 yr", "SOX, IRS"],
        ["Employee", "Personnel, reviews", "7yr post-term", "10yr post-term", "EEOC"],
        ["Contracts", "Agreements, SOWs", "Duration+6yr", "Duration+10yr", "Statute of limitations"],
        ["Customer PII", "Names, emails, payment", "Duration", "2yr post-end", "GDPR, CCPA"],
        ["App Logs", "Server, access, error", "90 days", "1 year", "Security"],
        ["Security Logs", "Auth, firewall, IDS", "1 year", "3 years", "PCI-DSS"],
        ["Board Minutes", "Minutes, resolutions", "Permanent", "Permanent", "Governance"],
        ["IP", "Patents, code", "Life+5yr", "Life+10yr", "IP law"],
    ])
    doc.add_heading("3. Legal Holds", level=1)
    doc.add_paragraph(
        "When litigation anticipated, Legal issues hold notice suspending schedules. "
        "Recipients must cease destruction. Violations: sanctions, adverse inference."
    )
    doc.add_heading("4. Disposal", level=1)
    _add_table(doc, ["Classification", "Electronic", "Paper"], [
        ["Public", "Standard deletion", "Recycling"],
        ["Internal", "Secure overwrite", "Cross-cut shredding"],
        ["Confidential", "Crypto erasure / DoD wipe", "Shredding + certificate"],
        ["Restricted", "Physical destruction", "Incineration + certificate"],
    ])
    doc.save(OUTPUT_DIR / "data_retention_policy.docx")
    print("  Created: data_retention_policy.docx")


def generate_code_of_conduct() -> None:
    """Doc 14: Code of Conduct."""
    doc = Document()
    doc.add_heading("Acme Corp Code of Conduct", level=0)
    doc.add_paragraph("Document ID: LEG-COC-001 | Version: 6.0 | Effective: January 1, 2026 | Owner: Legal + HR")
    doc.add_heading("1. Our Values", level=1)
    doc.add_paragraph(
        "Four core values: Customer First, Integrity Always, Continuous Improvement, "
        "Collaborative Excellence. This Code translates values into expected behaviors "
        "for all employees, contractors, and board members."
    )
    doc.add_heading("2. Ethical Business", level=1)
    doc.add_paragraph(
        "All business conducted honestly and ethically. No bribery, kickbacks, or "
        "corruption. Vendor gifts max $100, must be reported. No gifts to government "
        "officials. Annual anti-corruption training required."
    )
    doc.add_heading("3. Conflicts of Interest", level=1)
    doc.add_paragraph(
        "Avoid personal/company interest conflicts. Outside employment needs manager "
        "approval. Board seats at other companies need CEO approval. Financial interests "
        "in competitors/suppliers/customers must be disclosed. Manager-report romantic "
        "relationships must be disclosed to HR."
    )
    doc.add_heading("4. Workplace Respect", level=1)
    doc.add_paragraph(
        "Zero tolerance for harassment, discrimination, bullying. Protected characteristics: "
        "race, color, religion, sex, national origin, age, disability, sexual orientation, "
        "gender identity. Complaints investigated promptly and confidentially. "
        "No retaliation against reporters."
    )
    doc.add_heading("5. Confidentiality", level=1)
    doc.add_paragraph(
        "Protect confidential info during and after employment: trade secrets, customer "
        "data, pre-release financials, strategic plans, employee PII. No sharing with "
        "unauthorized persons, no public discussion, no social media posting."
    )
    doc.add_heading("6. Reporting", level=1)
    doc.add_paragraph(
        "Report violations via: manager, HR, Legal, ethics hotline (anonymous, "
        "1-800-555-ETHX), or ethics.acmecorp.com. All reports investigated. "
        "No retaliation against good-faith reporters."
    )
    doc.add_heading("7. Consequences", level=1)
    doc.add_paragraph(
        "Violations: disciplinary action up to termination. Serious violations may be "
        "reported to law enforcement. Annual acknowledgment required via compliance portal."
    )
    doc.save(OUTPUT_DIR / "code_of_conduct.docx")
    print("  Created: code_of_conduct.docx")


def generate_travel_safety() -> None:
    """Doc 15: Travel Safety Guidelines."""
    doc = Document()
    doc.add_heading("Acme Corp Travel Safety Guidelines", level=0)
    doc.add_paragraph("Document ID: SEC-TRV-001 | Version: 2.0 | Effective: January 1, 2026 | Owner: Security + HR")
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph("Ensure safety and security of employees during business travel.")
    doc.add_heading("2. Pre-Travel", level=1)
    doc.add_paragraph(
        "Register trip in Navan, check passport (6+ months validity), review travel "
        "advisories, get vaccinations, download safety app, share itinerary with manager."
    )
    doc.add_heading("3. Risk Levels", level=1)
    _add_table(doc, ["Level", "Destinations", "Approval", "Requirements"], [
        ["Low", "US, W. Europe, Japan, Australia", "Manager", "Standard procedures"],
        ["Medium", "E. Europe, most Asia, S. America", "VP + Security", "Security briefing, check-in"],
        ["High", "Conflict zones, high-crime", "CISO + CEO", "Security escort, sat phone, extraction plan"],
        ["Restricted", "Travel ban, active war zones", "Not permitted", "No travel authorized"],
    ])
    doc.add_heading("4. During Travel", level=1)
    doc.add_paragraph(
        "Keep devices locked and in possession. VPN for all internet. No public WiFi "
        "without VPN. No sensitive docs in hotel rooms. Use hotel safe. Be aware of "
        "surroundings, avoid high-risk areas after dark."
    )
    doc.add_heading("5. Device Security", level=1)
    doc.add_paragraph(
        "Medium/High destinations: clean travel laptop (no sensitive data), full disk "
        "encryption, disable Bluetooth/WiFi when unused, no public USB charging, "
        "report any tampering to IT Security on return."
    )
    doc.add_heading("6. Emergency", level=1)
    doc.add_paragraph(
        "24/7 hotline: +1-555-TRAVEL (+1-555-872-835). Contact local emergency services. "
        "Notify manager when safe. International SOS for medical evacuation. "
        "Policy: ACME-2026-INTL."
    )
    doc.add_heading("7. Post-Travel", level=1)
    doc.add_paragraph(
        "Medium/High returns: submit laptop to IT Security within 24 hours, change all "
        "passwords, report security concerns, complete post-travel debrief form."
    )
    doc.save(OUTPUT_DIR / "travel_safety_guidelines.docx")
    print("  Created: travel_safety_guidelines.docx")


def generate_all() -> None:
    generate_leave_benefits()
    generate_sdlc()
    generate_data_retention()
    generate_code_of_conduct()
    generate_travel_safety()
