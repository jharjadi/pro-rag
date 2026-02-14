#!/usr/bin/env python3
"""Generate 5 synthetic DOCX documents for the Phase 4a test corpus.

Topics:
1. IT Security Policy (with tables)
2. Employee Onboarding Guide (with tables)
3. Expense Reimbursement Policy (with tables)
4. Remote Work Policy
5. Incident Response Procedure (with tables)

Each document is ~2000-4000 words to produce meaningful chunks.
At least 3 documents contain tables to test table-aware chunking.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document


OUTPUT_DIR = Path(__file__).parent.parent / "data" / "test-corpus"


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value


# ---------------------------------------------------------------------------
# Document 1: IT Security Policy
# ---------------------------------------------------------------------------
def generate_it_security_policy() -> None:
    doc = Document()
    doc.add_heading("Acme Corp IT Security Policy", level=0)
    doc.add_paragraph(
        "Document ID: POL-SEC-001 | Version: 3.2 | Effective Date: January 1, 2026 | "
        "Classification: Internal | Owner: Chief Information Security Officer (CISO)"
    )

    doc.add_heading("1. Purpose and Scope", level=1)
    doc.add_paragraph(
        "This policy establishes the information security requirements for all Acme Corp "
        "employees, contractors, and third-party users who access Acme Corp information "
        "systems and data. The policy applies to all computing devices, networks, applications, "
        "and data repositories owned or managed by Acme Corp, including cloud-hosted services "
        "and bring-your-own-device (BYOD) equipment used for company business."
    )
    doc.add_paragraph(
        "The objectives of this policy are to: protect the confidentiality, integrity, and "
        "availability of Acme Corp information assets; ensure compliance with applicable laws, "
        "regulations, and contractual obligations; establish a framework for identifying and "
        "managing information security risks; and define roles and responsibilities for "
        "information security across the organization."
    )

    doc.add_heading("2. Data Classification", level=1)
    doc.add_paragraph(
        "All Acme Corp data must be classified according to the following scheme. Data owners "
        "are responsible for assigning the appropriate classification level at the time of "
        "creation. The classification determines the minimum security controls required for "
        "storage, transmission, and disposal."
    )
    _add_table(doc, ["Classification", "Description", "Examples", "Handling Requirements"], [
        ["Public", "Information approved for public release",
         "Marketing materials, press releases, public website content",
         "No restrictions on distribution. Standard backup procedures apply."],
        ["Internal", "General business information not intended for public release",
         "Internal memos, meeting notes, project plans, org charts",
         "Share within Acme Corp only. Encrypt in transit outside corporate network."],
        ["Confidential", "Sensitive business information that could cause harm if disclosed",
         "Financial reports, customer lists, strategic plans, source code",
         "Need-to-know access only. Encrypt at rest and in transit. Audit access logs quarterly."],
        ["Restricted", "Highly sensitive information subject to regulatory requirements",
         "PII, PHI, payment card data, trade secrets, encryption keys",
         "Strict access controls with MFA. Encrypt at rest (AES-256) and in transit (TLS 1.3). "
         "Log all access. Annual access review. Secure disposal required."],
    ])

    doc.add_heading("3. Access Control", level=1)
    doc.add_paragraph(
        "Access to Acme Corp information systems is granted based on the principle of least "
        "privilege. Users receive only the minimum access necessary to perform their job "
        "functions. All access requests must be approved by the user's manager and the system "
        "owner before provisioning."
    )

    doc.add_heading("3.1 Authentication Requirements", level=2)
    doc.add_paragraph(
        "All users must authenticate using their Acme Corp identity credentials. Multi-factor "
        "authentication (MFA) is mandatory for all remote access, privileged accounts, and "
        "access to Confidential or Restricted data. Single sign-on (SSO) via the corporate "
        "identity provider is the preferred authentication method for all applications."
    )
    _add_table(doc, ["Access Type", "Authentication Method", "MFA Required", "Session Timeout"], [
        ["Corporate network (on-site)", "SSO with corporate credentials",
         "No (network-level trust)", "8 hours"],
        ["VPN remote access", "SSO + MFA (authenticator app or hardware key)", "Yes", "12 hours"],
        ["Cloud applications (SaaS)", "SSO via SAML/OIDC",
         "Yes for Confidential+ data", "4 hours"],
        ["Privileged/admin access", "SSO + hardware security key (FIDO2)",
         "Yes (always)", "1 hour"],
        ["API/service accounts", "Certificate-based or OAuth2 client credentials",
         "N/A (non-interactive)", "Token-based expiry"],
    ])

    doc.add_heading("3.2 Password Policy", level=2)
    doc.add_paragraph(
        "Where passwords are used (legacy systems not yet integrated with SSO), the following "
        "requirements apply: minimum 14 characters; must include uppercase, lowercase, numbers, "
        "and special characters; passwords must not be reused within the last 24 password changes; "
        "passwords must be changed every 90 days for standard accounts and every 60 days for "
        "privileged accounts; account lockout after 5 failed attempts with 30-minute lockout period."
    )

    doc.add_heading("4. Network Security", level=1)
    doc.add_paragraph(
        "The Acme Corp network is segmented into security zones to limit the blast radius of "
        "potential security incidents. Network traffic between zones is controlled by firewalls "
        "with explicit allow rules. All inter-zone traffic is logged and monitored."
    )
    _add_table(doc, ["Zone", "Purpose", "Access Restrictions", "Monitoring Level"], [
        ["DMZ", "Public-facing services (web servers, load balancers)",
         "Inbound: HTTPS only. Outbound: restricted to application tier.",
         "Real-time IDS/IPS"],
        ["Application Tier", "Internal application servers and APIs",
         "Inbound: from DMZ and corporate network only. No direct internet access.",
         "Real-time IDS, daily log review"],
        ["Data Tier", "Databases, file servers, backup systems",
         "Inbound: from application tier only. No direct user access.",
         "Real-time IDS/IPS, continuous audit logging"],
        ["Corporate Network", "Employee workstations, printers, internal tools",
         "Outbound: filtered through web proxy. Inbound: VPN only.",
         "Network flow analysis, endpoint detection"],
        ["Management Network", "Infrastructure management (IPMI, switches, firewalls)",
         "Isolated. Access via jump host with MFA only.",
         "Full packet capture, real-time alerting"],
    ])

    doc.add_heading("5. Endpoint Security", level=1)
    doc.add_paragraph(
        "All Acme Corp managed endpoints must have the following security controls installed "
        "and active: endpoint detection and response (EDR) agent with real-time threat detection; "
        "full-disk encryption (FileVault for macOS, BitLocker for Windows); automatic OS and "
        "application patching within 14 days of release for standard patches and 48 hours for "
        "critical security patches; host-based firewall enabled with default-deny inbound policy; "
        "USB mass storage devices disabled by default (exceptions require CISO approval)."
    )
    doc.add_paragraph(
        "BYOD devices used for company business must be enrolled in the mobile device management "
        "(MDM) system and comply with the following minimum requirements: device encryption enabled, "
        "screen lock with biometric or 6-digit PIN, remote wipe capability, and separation of "
        "corporate and personal data via containerization."
    )

    doc.add_heading("6. Incident Reporting", level=1)
    doc.add_paragraph(
        "All employees must report suspected security incidents immediately to the Security "
        "Operations Center (SOC) via email at security@acmecorp.com or by calling the 24/7 "
        "security hotline at ext. 9-1-1. Incidents include but are not limited to: suspected "
        "phishing emails, unauthorized access attempts, lost or stolen devices, malware "
        "infections, data breaches, and social engineering attempts."
    )
    doc.add_paragraph(
        "The SOC will triage all reported incidents within 15 minutes during business hours "
        "and within 1 hour outside business hours. Critical incidents (data breach, active "
        "intrusion, ransomware) trigger the Incident Response Plan (see POL-IR-001) with "
        "automatic escalation to the CISO and legal counsel."
    )

    doc.add_heading("7. Compliance and Audit", level=1)
    doc.add_paragraph(
        "Compliance with this policy is mandatory. Violations may result in disciplinary action "
        "up to and including termination of employment or contract. The Information Security "
        "team conducts quarterly access reviews, annual penetration testing, and continuous "
        "vulnerability scanning. External audits are conducted annually by an independent "
        "third-party assessor. All audit findings must be remediated within the timeframes "
        "specified in the risk treatment plan."
    )

    doc.add_heading("8. Policy Review", level=1)
    doc.add_paragraph(
        "This policy is reviewed annually by the CISO and updated as necessary to reflect "
        "changes in the threat landscape, regulatory requirements, or business operations. "
        "All employees must acknowledge receipt and understanding of this policy annually "
        "through the compliance training portal."
    )

    doc.save(OUTPUT_DIR / "it_security_policy.docx")
    print("  Created: it_security_policy.docx")


# ---------------------------------------------------------------------------
# Document 2: Employee Onboarding Guide
# ---------------------------------------------------------------------------
def generate_onboarding_guide() -> None:
    doc = Document()
    doc.add_heading("Acme Corp Employee Onboarding Guide", level=0)
    doc.add_paragraph(
        "Document ID: HR-ONB-001 | Version: 2.1 | Last Updated: December 15, 2025 | "
        "Owner: Human Resources Department"
    )

    doc.add_heading("1. Welcome to Acme Corp", level=1)
    doc.add_paragraph(
        "Welcome to Acme Corp! We are excited to have you join our team. This guide will walk "
        "you through everything you need to know during your first 90 days. Our onboarding "
        "program is designed to help you integrate smoothly into the company, understand our "
        "culture and values, and set you up for success in your new role."
    )
    doc.add_paragraph(
        "Acme Corp was founded in 2010 with a mission to deliver innovative enterprise solutions "
        "that help businesses operate more efficiently. Today, we serve over 2,000 customers "
        "across 15 countries with a team of 850 employees. Our core values are: Customer First, "
        "Integrity Always, Continuous Improvement, and Collaborative Excellence."
    )

    doc.add_heading("2. First Day Checklist", level=1)
    doc.add_paragraph(
        "Your first day is primarily about getting set up and meeting your team. Please arrive "
        "at the office by 9:00 AM (or log in to the virtual onboarding session if remote). "
        "Your manager and an onboarding buddy will guide you through the day."
    )
    _add_table(doc, ["Time", "Activity", "Location/Link", "Contact"], [
        ["9:00 AM", "Welcome and badge/laptop pickup", "Reception Desk, Building A",
         "Facilities Team (facilities@acmecorp.com)"],
        ["9:30 AM", "IT setup: laptop configuration, accounts, VPN", "IT Help Desk, Room 102",
         "IT Support (itsupport@acmecorp.com)"],
        ["10:30 AM", "HR orientation: benefits enrollment, policies overview",
         "HR Conference Room / Zoom", "HR Onboarding (onboarding@acmecorp.com)"],
        ["12:00 PM", "Team lunch with your manager and onboarding buddy",
         "TBD (manager will arrange)", "Your Manager"],
        ["1:30 PM", "Team introduction and workspace tour", "Your team's area",
         "Onboarding Buddy"],
        ["3:00 PM", "Security awareness training (mandatory)", "Online via LMS",
         "Security Team (security@acmecorp.com)"],
        ["4:00 PM", "1:1 with your manager: role expectations, 30-60-90 day goals",
         "Manager's office / Zoom", "Your Manager"],
    ])

    doc.add_heading("3. Systems and Tools Access", level=1)
    doc.add_paragraph(
        "You will receive access to the following systems during your first week. All access "
        "is provisioned through our identity management system and requires SSO authentication. "
        "If you need access to additional systems, submit a request through the IT Service Portal."
    )
    _add_table(doc, ["System", "Purpose", "Access Level", "Provisioning Timeline"], [
        ["Google Workspace", "Email, calendar, documents, video conferencing",
         "Full access", "Day 1"],
        ["Slack", "Team communication and channels", "Full access", "Day 1"],
        ["Jira", "Project management and issue tracking", "Team-level access", "Day 1-2"],
        ["GitHub", "Source code repositories", "Team repositories (read/write)", "Day 1-2"],
        ["Confluence", "Knowledge base and documentation",
         "Company-wide read, team-level write", "Day 1"],
        ["Workday", "HR self-service (pay stubs, time off, benefits)",
         "Self-service access", "Day 1"],
        ["Salesforce", "CRM (sales and customer success teams only)", "Role-based",
         "Week 1 (if applicable)"],
        ["AWS Console", "Cloud infrastructure (engineering teams only)",
         "Read-only initially", "Week 1-2 (if applicable)"],
    ])

    doc.add_heading("4. Benefits Overview", level=1)
    doc.add_paragraph(
        "Acme Corp offers a comprehensive benefits package that begins on your first day of "
        "employment. You must complete benefits enrollment within 30 days of your start date "
        "through the Workday self-service portal. If you do not enroll within 30 days, you "
        "will be automatically enrolled in the default plan options."
    )

    doc.add_heading("4.1 Health and Wellness", level=2)
    doc.add_paragraph(
        "We offer three medical plan options through Blue Cross Blue Shield: a PPO plan with "
        "broad network access, an HMO plan with lower premiums, and a High Deductible Health "
        "Plan (HDHP) paired with a Health Savings Account (HSA). Acme Corp contributes $1,500 "
        "annually to your HSA if you choose the HDHP option. Dental and vision coverage is "
        "included in all medical plans at no additional cost."
    )
    doc.add_paragraph(
        "Additional wellness benefits include: $500 annual wellness stipend for gym memberships, "
        "fitness equipment, or wellness apps; free access to the Employee Assistance Program (EAP) "
        "for mental health counseling (up to 12 sessions per year); annual health screening and "
        "flu vaccination at no cost; and ergonomic workstation assessment for both office and "
        "remote workers."
    )

    doc.add_heading("4.2 Time Off", level=2)
    doc.add_paragraph(
        "Acme Corp provides generous paid time off to support work-life balance. PTO accrues "
        "from your first day and can be used after 30 days of employment. Unused PTO carries "
        "over up to a maximum of 5 days per year."
    )
    _add_table(doc, ["Benefit", "Allowance", "Accrual", "Notes"], [
        ["Paid Time Off (PTO)", "20 days/year (0-3 years tenure)", "1.67 days/month",
         "Increases to 25 days after 3 years, 30 days after 7 years"],
        ["Sick Leave", "10 days/year", "Granted January 1", "Does not carry over"],
        ["Public Holidays", "11 days/year", "Fixed schedule", "See company calendar"],
        ["Parental Leave", "16 weeks (primary), 8 weeks (secondary)", "N/A",
         "Available after 6 months of employment"],
        ["Bereavement Leave", "5 days (immediate family), 3 days (extended)", "As needed",
         "Immediate family: spouse, parent, child, sibling"],
        ["Volunteer Day", "2 days/year", "N/A",
         "Pre-approved community service activities"],
    ])

    doc.add_heading("5. Training and Development", level=1)
    doc.add_paragraph(
        "Acme Corp invests in your professional growth. Every employee has access to a $3,000 "
        "annual learning and development budget that can be used for conferences, courses, "
        "certifications, and books. Additionally, we provide internal training programs including "
        "leadership development, technical skills workshops, and cross-functional rotations."
    )
    doc.add_paragraph(
        "During your first 90 days, you are expected to complete the following mandatory training "
        "modules through our Learning Management System (LMS): Security Awareness Training "
        "(due by end of Week 1), Code of Conduct and Ethics (due by end of Week 2), Data Privacy "
        "and GDPR Fundamentals (due by end of Week 3), and Diversity, Equity, and Inclusion "
        "Foundations (due by end of Month 1)."
    )

    doc.add_heading("6. Performance and Feedback", level=1)
    doc.add_paragraph(
        "Your performance will be evaluated through regular check-ins and formal reviews. During "
        "your first 90 days, you will have weekly 1:1 meetings with your manager to discuss "
        "progress, challenges, and goals. After the initial period, 1:1s move to bi-weekly. "
        "Formal performance reviews are conducted semi-annually in June and December."
    )
    doc.add_paragraph(
        "We use a continuous feedback model. You are encouraged to give and receive feedback "
        "regularly through our peer feedback tool in Workday. At the 30-day, 60-day, and 90-day "
        "marks, your manager will conduct a structured check-in to assess your onboarding "
        "progress and address any concerns."
    )

    doc.add_heading("7. Key Contacts", level=1)
    doc.add_paragraph(
        "If you have questions during your onboarding, don't hesitate to reach out to any of "
        "the following contacts. Your onboarding buddy is your first point of contact for "
        "day-to-day questions about how things work at Acme Corp."
    )
    _add_table(doc, ["Department", "Contact", "Email", "For Help With"], [
        ["HR Onboarding", "Sarah Chen", "onboarding@acmecorp.com",
         "Benefits, policies, general HR questions"],
        ["IT Support", "Help Desk", "itsupport@acmecorp.com",
         "Laptop, accounts, VPN, software access"],
        ["Facilities", "Mike Johnson", "facilities@acmecorp.com",
         "Badge, parking, office supplies, workspace"],
        ["Security", "Security Team", "security@acmecorp.com",
         "Security training, incident reporting"],
        ["Payroll", "Payroll Team", "payroll@acmecorp.com",
         "Pay stubs, tax forms, direct deposit"],
    ])

    doc.save(OUTPUT_DIR / "employee_onboarding_guide.docx")
    print("  Created: employee_onboarding_guide.docx")


# ---------------------------------------------------------------------------
# Document 3: Expense Reimbursement Policy
# ---------------------------------------------------------------------------
def generate_expense_policy() -> None:
    doc = Document()
    doc.add_heading("Acme Corp Expense Reimbursement Policy", level=0)
    doc.add_paragraph(
        "Document ID: FIN-EXP-001 | Version: 4.0 | Effective Date: January 1, 2026 | "
        "Owner: Finance Department | Approved by: CFO"
    )

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This policy establishes guidelines for the reimbursement of business expenses incurred "
        "by Acme Corp employees in the course of performing their job duties. The policy ensures "
        "that expenses are reasonable, properly documented, and approved in a timely manner. "
        "All employees are expected to exercise good judgment and fiscal responsibility when "
        "incurring business expenses."
    )

    doc.add_heading("2. Scope", level=1)
    doc.add_paragraph(
        "This policy applies to all full-time and part-time employees of Acme Corp. Contractors "
        "and consultants should refer to their contract terms for expense reimbursement provisions. "
        "Interns are eligible for expense reimbursement only with prior written approval from "
        "their manager and the Finance department."
    )

    doc.add_heading("3. Expense Categories and Limits", level=1)
    doc.add_paragraph(
        "The following table outlines the approved expense categories and their associated limits. "
        "Expenses exceeding these limits require pre-approval from the employee's VP-level manager "
        "and the Finance department. Any expense over $5,000 requires CFO approval regardless "
        "of category."
    )
    _add_table(doc,
        ["Category", "Per-Occurrence Limit", "Monthly Limit", "Approval Required", "Receipt Required"],
        [
            ["Domestic airfare", "$800 (economy class)", "No monthly limit", "Manager", "Yes"],
            ["International airfare", "$3,000 (economy/premium economy)", "No monthly limit",
             "VP + Finance", "Yes"],
            ["Hotel/lodging", "$250/night (domestic), $350/night (international)",
             "No monthly limit", "Manager", "Yes"],
            ["Meals (solo, while traveling)", "$75/day", "$750/month", "Manager",
             "Yes (over $25)"],
            ["Client entertainment meals", "$150/person", "$1,000/month", "Manager",
             "Yes (with attendee list)"],
            ["Ground transportation", "$100/day", "$500/month", "Manager", "Yes (over $25)"],
            ["Rental car", "$75/day (compact/midsize)", "No monthly limit", "Manager", "Yes"],
            ["Mileage (personal vehicle)", "$0.67/mile (IRS rate)", "$500/month", "Manager",
             "Mileage log"],
            ["Office supplies", "$100/occurrence", "$200/month", "Self-approved", "Yes"],
            ["Software/subscriptions", "$50/month per tool", "$200/month total",
             "Manager + IT", "Yes"],
            ["Professional development", "$3,000/year", "N/A (annual)", "Manager", "Yes"],
            ["Home office equipment", "$500/year", "N/A (annual)", "Manager + IT", "Yes"],
        ],
    )

    doc.add_heading("4. Travel Booking Procedures", level=1)
    doc.add_paragraph(
        "All business travel must be booked through the Acme Corp travel portal (powered by "
        "Navan/TripActions) to ensure compliance with negotiated corporate rates and travel "
        "policies. Bookings made outside the travel portal will only be reimbursed if the "
        "employee can demonstrate that the portal was unavailable or that the external booking "
        "resulted in a lower cost."
    )

    doc.add_heading("4.1 Air Travel", level=2)
    doc.add_paragraph(
        "Economy class is the standard for all domestic flights and international flights under "
        "6 hours. Premium economy is permitted for international flights over 6 hours. Business "
        "class requires VP-level pre-approval and is generally reserved for flights over 10 hours "
        "or when required for medical reasons. First class is not reimbursable under any "
        "circumstances."
    )
    doc.add_paragraph(
        "Employees should book flights at least 14 days in advance when possible to secure "
        "lower fares. Same-day or next-day bookings require manager approval. Frequent flyer "
        "miles earned on business travel belong to the employee but should not influence booking "
        "decisions (i.e., do not choose a more expensive flight to earn miles on a preferred airline)."
    )

    doc.add_heading("4.2 Hotel and Lodging", level=2)
    doc.add_paragraph(
        "Hotel bookings should be made through the travel portal, which provides access to "
        "negotiated corporate rates at preferred hotel chains. The nightly rate limits ($250 "
        "domestic, $350 international) include the room rate and applicable taxes but exclude "
        "incidental charges. Employees traveling to high-cost cities (New York, San Francisco, "
        "London, Tokyo, Sydney) may request a rate exception of up to 150% of the standard limit "
        "with manager approval."
    )

    doc.add_heading("5. Expense Report Submission", level=1)
    doc.add_paragraph(
        "Expense reports must be submitted within 30 days of incurring the expense or within "
        "10 business days of returning from a trip, whichever comes first. Late submissions "
        "(31-60 days) require manager and Finance approval. Submissions older than 60 days "
        "will not be reimbursed except in extraordinary circumstances approved by the CFO."
    )
    doc.add_paragraph(
        "All expense reports are submitted through the Concur expense management system. Each "
        "report must include: a clear business purpose for each expense, itemized receipts for "
        "all expenses over $25, the names and business relationship of all attendees for "
        "entertainment expenses, and the project or cost center code to charge."
    )

    doc.add_heading("5.1 Approval Workflow", level=2)
    _add_table(doc, ["Expense Amount", "Approval Chain", "Expected Turnaround"], [
        ["Under $500", "Direct manager", "3 business days"],
        ["$500 - $2,000", "Direct manager then Department head", "5 business days"],
        ["$2,000 - $5,000", "Direct manager then VP", "7 business days"],
        ["Over $5,000", "Direct manager then VP then CFO", "10 business days"],
    ])

    doc.add_heading("6. Non-Reimbursable Expenses", level=1)
    doc.add_paragraph(
        "The following expenses are not eligible for reimbursement under any circumstances: "
        "personal entertainment or recreation; alcoholic beverages (except as part of approved "
        "client entertainment with a $50/person limit); traffic violations, parking tickets, or "
        "towing charges; airline upgrades not pre-approved; personal travel insurance; laundry "
        "and dry cleaning (unless trip exceeds 5 consecutive days); spouse or family member "
        "travel expenses; political contributions or charitable donations; and any expense "
        "that violates Acme Corp's Code of Conduct."
    )

    doc.add_heading("7. Corporate Credit Card", level=1)
    doc.add_paragraph(
        "Employees who travel frequently (more than 4 trips per year) or regularly incur "
        "business expenses may apply for an Acme Corp corporate credit card. The card is issued "
        "in the employee's name but is the property of Acme Corp. Personal charges on the "
        "corporate card are strictly prohibited and may result in disciplinary action."
    )
    doc.add_paragraph(
        "Corporate card statements are automatically imported into Concur. Employees must "
        "reconcile and submit expense reports for all corporate card charges within 30 days "
        "of the statement date. Unreconciled charges older than 60 days will be deducted from "
        "the employee's paycheck after written notice."
    )

    doc.add_heading("8. Audit and Compliance", level=1)
    doc.add_paragraph(
        "The Finance department conducts random audits of expense reports on a quarterly basis. "
        "Approximately 10% of all submitted reports are selected for detailed review. Employees "
        "must retain original receipts for 3 years and make them available upon request. "
        "Fraudulent expense claims will result in immediate termination and may be referred "
        "to law enforcement."
    )

    doc.save(OUTPUT_DIR / "expense_reimbursement_policy.docx")
    print("  Created: expense_reimbursement_policy.docx")


# ---------------------------------------------------------------------------
# Document 4: Remote Work Policy
# ---------------------------------------------------------------------------
def generate_remote_work_policy() -> None:
    doc = Document()
    doc.add_heading("Acme Corp Remote Work Policy", level=0)
    doc.add_paragraph(
        "Document ID: HR-RW-001 | Version: 2.0 | Effective Date: March 1, 2026 | "
        "Owner: Human Resources Department | Approved by: Chief People Officer"
    )

    doc.add_heading("1. Purpose and Philosophy", level=1)
    doc.add_paragraph(
        "Acme Corp recognizes that flexible work arrangements contribute to employee satisfaction, "
        "productivity, and talent retention. This policy establishes the framework for remote "
        "and hybrid work arrangements, ensuring that all employees can perform their duties "
        "effectively regardless of their physical location while maintaining the collaboration "
        "and culture that define Acme Corp."
    )
    doc.add_paragraph(
        "Our approach to remote work is built on trust and accountability. We believe that "
        "results matter more than physical presence, and we empower our employees to manage "
        "their time and workspace in ways that optimize their productivity. At the same time, "
        "we recognize that in-person collaboration has unique value for building relationships, "
        "brainstorming, and complex problem-solving."
    )

    doc.add_heading("2. Eligibility", level=1)
    doc.add_paragraph(
        "All full-time employees who have completed their 90-day probationary period are "
        "eligible to request a remote or hybrid work arrangement. Eligibility is determined "
        "by the employee's role requirements, performance history, and manager approval. Some "
        "roles may require regular on-site presence due to the nature of the work (e.g., "
        "facilities management, hardware engineering, reception)."
    )
    doc.add_paragraph(
        "Employees on a Performance Improvement Plan (PIP) are not eligible for remote work "
        "until the PIP is successfully completed. New employees during their probationary period "
        "are expected to work on-site or in a hybrid arrangement (minimum 3 days in office) to "
        "facilitate onboarding and team integration."
    )

    doc.add_heading("3. Work Arrangement Types", level=1)
    doc.add_paragraph(
        "Acme Corp offers three work arrangement types. The default arrangement is hybrid, "
        "which balances remote flexibility with in-person collaboration. Employees may request "
        "a different arrangement through their manager, subject to role requirements "
        "and business needs."
    )

    doc.add_heading("3.1 On-Site", level=2)
    doc.add_paragraph(
        "On-site employees work from an Acme Corp office five days per week. This arrangement "
        "is required for roles that involve physical equipment, secure facilities, or front-desk "
        "responsibilities. On-site employees receive a $50/month commuter benefit for public "
        "transit or parking."
    )

    doc.add_heading("3.2 Hybrid", level=2)
    doc.add_paragraph(
        "Hybrid employees split their time between the office and a remote location. The standard "
        "hybrid schedule is 2-3 days in the office per week, with specific days determined by "
        "the team's collaboration needs. Teams are encouraged to designate anchor days when "
        "all team members are in the office for meetings, planning sessions, and social events. "
        "Hybrid employees must live within reasonable commuting distance of their assigned office "
        "(generally within 75 miles or 90 minutes)."
    )

    doc.add_heading("3.3 Fully Remote", level=2)
    doc.add_paragraph(
        "Fully remote employees work from a location of their choice and are not expected to "
        "come to the office on a regular basis. This arrangement is available for roles that "
        "can be performed entirely remotely and requires VP-level approval. Fully remote employees "
        "are expected to travel to the office quarterly for team events and planning sessions "
        "(travel expenses covered by the company). Fully remote employees must reside in a state "
        "or country where Acme Corp has an established legal entity for tax and employment law "
        "compliance."
    )

    doc.add_heading("4. Home Office Requirements", level=1)
    doc.add_paragraph(
        "Employees working remotely (hybrid or fully remote) must maintain a dedicated workspace "
        "that meets the following requirements: a quiet, private area suitable for video calls "
        "and focused work; reliable high-speed internet connection (minimum 50 Mbps download, "
        "10 Mbps upload); a desk and ergonomic chair (Acme Corp provides a $500 home office "
        "stipend for initial setup); and adequate lighting for video conferencing."
    )
    doc.add_paragraph(
        "Remote employees are responsible for the security of their home workspace. This includes "
        "locking their computer when stepping away, using the company VPN for all work activities, "
        "ensuring that family members or roommates cannot access company data, and following all "
        "requirements in the IT Security Policy (POL-SEC-001). Acme Corp reserves the right to "
        "conduct virtual workspace assessments to verify compliance."
    )

    doc.add_heading("5. Communication and Availability", level=1)
    doc.add_paragraph(
        "Remote and hybrid employees must be available during core business hours (10:00 AM to "
        "3:00 PM in their local time zone) for meetings, collaboration, and urgent requests. "
        "Outside of core hours, employees have flexibility to structure their workday as they "
        "see fit, provided they meet their performance objectives and total weekly hours."
    )
    doc.add_paragraph(
        "All remote employees must keep their Slack status updated to reflect their availability. "
        "Video should be enabled by default for team meetings and 1:1s to maintain personal "
        "connection. Response time expectations: Slack messages within 1 hour during core hours, "
        "emails within 4 hours during business hours, and urgent matters via phone call."
    )

    doc.add_heading("6. Equipment and Technology", level=1)
    doc.add_paragraph(
        "Acme Corp provides all employees with a company laptop, monitor, keyboard, and mouse. "
        "Remote employees may request additional equipment through the IT Service Portal. The "
        "company provides a $100/month stipend for internet and phone expenses for fully remote "
        "employees and a $50/month stipend for hybrid employees. All company equipment must be "
        "returned upon termination of employment or change to on-site arrangement."
    )

    doc.add_heading("7. Performance Management", level=1)
    doc.add_paragraph(
        "Remote work arrangements are contingent on maintaining satisfactory performance. "
        "Managers are expected to evaluate remote employees based on outcomes and deliverables, "
        "not hours logged or online presence. If an employee's performance declines after "
        "transitioning to remote work, the manager may require a return to on-site or hybrid "
        "arrangement after consultation with HR."
    )
    doc.add_paragraph(
        "Remote employees participate in the same performance review cycle as on-site employees. "
        "Managers should schedule regular 1:1 check-ins (at least bi-weekly) and ensure remote "
        "employees have equal access to career development opportunities, promotions, and "
        "high-visibility projects. The company monitors for proximity bias in performance "
        "ratings and promotion decisions."
    )

    doc.add_heading("8. Tax and Legal Considerations", level=1)
    doc.add_paragraph(
        "Employees are responsible for understanding the tax implications of their work location. "
        "Working from a different state or country may create tax obligations for both the "
        "employee and Acme Corp. Employees must notify HR before working from a location outside "
        "their approved work state/country for more than 14 consecutive days. International "
        "remote work requires pre-approval from HR and Legal due to immigration, tax, and "
        "employment law complexities."
    )

    doc.add_heading("9. Policy Modifications and Revocation", level=1)
    doc.add_paragraph(
        "Acme Corp reserves the right to modify or revoke remote work arrangements at any time "
        "based on business needs, performance concerns, or changes in company policy. Employees "
        "will receive at least 30 days notice before any mandatory change to their work "
        "arrangement, except in cases of performance issues or security concerns that require "
        "immediate action. This policy is reviewed annually by the HR department and updated "
        "as needed."
    )

    doc.save(OUTPUT_DIR / "remote_work_policy.docx")
    print("  Created: remote_work_policy.docx")


# ---------------------------------------------------------------------------
# Document 5: Incident Response Procedure
# ---------------------------------------------------------------------------
def generate_incident_response() -> None:
    doc = Document()
    doc.add_heading("Acme Corp Incident Response Procedure", level=0)
    doc.add_paragraph(
        "Document ID: POL-IR-001 | Version: 3.0 | Effective Date: January 15, 2026 | "
        "Classification: Confidential | Owner: Chief Information Security Officer (CISO)"
    )

    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This document defines the procedures for detecting, responding to, containing, and "
        "recovering from information security incidents at Acme Corp. The goal is to minimize "
        "the impact of security incidents on business operations, protect sensitive data, and "
        "ensure compliance with legal and regulatory notification requirements."
    )
    doc.add_paragraph(
        "This procedure is activated when a security event is confirmed as an incident by the "
        "Security Operations Center (SOC) or when a critical alert is triggered by automated "
        "monitoring systems. All employees are required to be familiar with the incident "
        "reporting process described in Section 3."
    )

    doc.add_heading("2. Incident Classification", level=1)
    doc.add_paragraph(
        "Security incidents are classified by severity level, which determines the response "
        "urgency, escalation path, and communication requirements. The SOC assigns the initial "
        "severity level based on the criteria below. Severity may be upgraded or downgraded "
        "as more information becomes available during the investigation."
    )
    _add_table(doc,
        ["Severity", "Name", "Description", "Response Time", "Examples"],
        [
            ["SEV-1", "Critical",
             "Active data breach, ransomware, or complete service outage affecting customers",
             "Immediate (within 15 minutes)", "Active data exfiltration, ransomware encryption "
             "in progress, complete production outage, compromised admin credentials"],
            ["SEV-2", "High",
             "Confirmed compromise with potential for data loss or significant service degradation",
             "Within 1 hour", "Malware on server, unauthorized access to Confidential data, "
             "partial service outage, compromised user account with elevated privileges"],
            ["SEV-3", "Medium",
             "Suspicious activity requiring investigation, limited impact",
             "Within 4 hours", "Phishing email clicked (no credential entry), unusual login "
             "patterns, vulnerability discovered in production, failed brute-force attempts"],
            ["SEV-4", "Low",
             "Minor security event, no confirmed compromise",
             "Within 24 hours", "Phishing email reported (not clicked), policy violation "
             "(e.g., unauthorized software), lost badge, minor vulnerability in non-production"],
        ],
    )

    doc.add_heading("3. Incident Reporting", level=1)
    doc.add_paragraph(
        "Any employee who suspects a security incident must report it immediately through one "
        "of the following channels: email to security@acmecorp.com (monitored 24/7), call the "
        "security hotline at extension 9-1-1, submit a ticket through the IT Service Portal "
        "under 'Security Incident', or notify their manager who will escalate to the SOC."
    )
    doc.add_paragraph(
        "When reporting an incident, provide as much of the following information as possible: "
        "what happened (description of the event), when it happened (date and time), how it was "
        "discovered, what systems or data are affected, what actions have been taken so far, and "
        "contact information for follow-up questions. Do not attempt to investigate or remediate "
        "the incident yourself unless you are a member of the Incident Response Team."
    )

    doc.add_heading("4. Incident Response Team", level=1)
    doc.add_paragraph(
        "The Incident Response Team (IRT) is a cross-functional team responsible for managing "
        "security incidents from detection through resolution. The team is led by the Incident "
        "Commander (IC), who has authority to make decisions and allocate resources during an "
        "active incident."
    )
    _add_table(doc, ["Role", "Responsibility", "Primary", "Backup"], [
        ["Incident Commander", "Overall incident management, decision-making, communication",
         "CISO (Alex Rivera)", "Director of Security (Jordan Lee)"],
        ["Technical Lead", "Technical investigation, containment, and remediation",
         "Senior Security Engineer", "On-call security engineer"],
        ["Communications Lead", "Internal and external communications, stakeholder updates",
         "VP of Communications", "HR Director"],
        ["Legal Advisor", "Legal obligations, regulatory notifications, evidence preservation",
         "General Counsel", "Outside counsel (on retainer)"],
        ["Business Liaison", "Business impact assessment, customer communication",
         "VP of affected business unit", "COO"],
    ])

    doc.add_heading("5. Response Phases", level=1)

    doc.add_heading("5.1 Phase 1: Detection and Triage", level=2)
    doc.add_paragraph(
        "Upon receiving an incident report or automated alert, the SOC analyst on duty performs "
        "initial triage: verify the report (is this a real incident or a false positive?), "
        "assign initial severity level, create an incident ticket in the incident management "
        "system, and notify the Incident Commander if severity is SEV-1 or SEV-2. For SEV-1 "
        "incidents, the IC is notified immediately via phone call and a war room (virtual or "
        "physical) is established within 15 minutes."
    )

    doc.add_heading("5.2 Phase 2: Containment", level=2)
    doc.add_paragraph(
        "The goal of containment is to limit the scope and impact of the incident while "
        "preserving evidence for investigation. Containment strategies depend on the type of "
        "incident. Short-term containment actions may include: isolating affected systems from "
        "the network, blocking malicious IP addresses or domains at the firewall, disabling "
        "compromised user accounts, revoking API keys or access tokens, and redirecting traffic "
        "away from affected services."
    )
    doc.add_paragraph(
        "Long-term containment involves implementing temporary fixes that allow business "
        "operations to continue while the root cause is being addressed. This may include "
        "deploying additional monitoring, implementing temporary access restrictions, or "
        "activating backup systems. All containment actions must be documented in the incident "
        "ticket with timestamps."
    )

    doc.add_heading("5.3 Phase 3: Eradication", level=2)
    doc.add_paragraph(
        "Once the incident is contained, the Technical Lead coordinates eradication activities "
        "to remove the threat from the environment. This includes: identifying and removing "
        "malware or unauthorized access mechanisms, patching exploited vulnerabilities, resetting "
        "compromised credentials (all affected accounts, not just the initially identified ones), "
        "reviewing and hardening configurations that contributed to the incident, and scanning "
        "for indicators of compromise (IOCs) across the environment to ensure the threat is "
        "fully eliminated."
    )

    doc.add_heading("5.4 Phase 4: Recovery", level=2)
    doc.add_paragraph(
        "Recovery involves restoring affected systems to normal operation. The Technical Lead "
        "develops a recovery plan that includes: restoring systems from known-good backups if "
        "necessary, gradually bringing systems back online with enhanced monitoring, verifying "
        "system integrity before returning to production, monitoring for signs of re-compromise "
        "for at least 72 hours after recovery, and confirming with business stakeholders that "
        "services are functioning normally."
    )

    doc.add_heading("5.5 Phase 5: Post-Incident Review", level=2)
    doc.add_paragraph(
        "A post-incident review (PIR) is conducted within 5 business days of incident closure "
        "for all SEV-1 and SEV-2 incidents, and within 10 business days for SEV-3 incidents. "
        "The PIR is a blameless review focused on learning and improvement. The review covers: "
        "timeline of events (what happened and when), root cause analysis, effectiveness of "
        "the response (what worked well, what could be improved), action items for preventing "
        "similar incidents, and updates to runbooks, monitoring, or procedures."
    )
    doc.add_paragraph(
        "The PIR report is documented in the incident management system and shared with "
        "relevant stakeholders. Action items are tracked to completion with assigned owners "
        "and due dates. The CISO reviews all PIR reports monthly and presents a summary to "
        "the executive team quarterly."
    )

    doc.add_heading("6. Communication Protocols", level=1)
    doc.add_paragraph(
        "Clear and timely communication is critical during incident response. The Communications "
        "Lead manages all internal and external communications according to the following "
        "guidelines."
    )
    _add_table(doc,
        ["Severity", "Internal Communication", "Customer Communication", "Regulatory Notification"],
        [
            ["SEV-1", "Immediate: exec team, affected teams, all-hands if widespread. "
             "Updates every 30 minutes until contained.",
             "Within 4 hours of confirmation. Status page updated. Direct notification to "
             "affected customers.",
             "Within 72 hours per GDPR. Within 24 hours per state breach notification laws "
             "(varies by jurisdiction)."],
            ["SEV-2", "Within 1 hour: exec team, affected teams. Updates every 2 hours.",
             "Within 24 hours if customer data affected. Status page updated if service impacted.",
             "As required based on data types involved."],
            ["SEV-3", "Within 4 hours: security team, affected team leads. Daily updates.",
             "Only if customer-facing impact confirmed.",
             "Generally not required unless PII involved."],
            ["SEV-4", "Documented in incident ticket. Weekly summary to security team.",
             "Not required.", "Not required."],
        ],
    )

    doc.add_heading("7. Evidence Preservation", level=1)
    doc.add_paragraph(
        "Preserving evidence is critical for investigation, legal proceedings, and regulatory "
        "compliance. The following evidence preservation procedures must be followed for all "
        "SEV-1 and SEV-2 incidents: capture full disk images of affected systems before any "
        "remediation, preserve all relevant log files (system, application, network, authentication), "
        "document the chain of custody for all evidence, store evidence in the designated secure "
        "evidence repository with access restricted to the IRT, and retain evidence for a minimum "
        "of 7 years or as required by applicable regulations."
    )

    doc.add_heading("8. Training and Exercises", level=1)
    doc.add_paragraph(
        "All members of the Incident Response Team must complete annual incident response training "
        "and participate in at least two tabletop exercises per year. Tabletop exercises simulate "
        "realistic incident scenarios and test the team's ability to follow procedures, make "
        "decisions under pressure, and communicate effectively. Exercise findings are incorporated "
        "into procedure updates and training improvements."
    )
    doc.add_paragraph(
        "Additionally, the SOC conducts quarterly red team exercises to test detection and "
        "response capabilities. These exercises are coordinated with the CISO and may involve "
        "simulated phishing campaigns, penetration testing, or simulated data exfiltration "
        "attempts. Results are documented and used to improve monitoring rules, response "
        "procedures, and employee awareness training."
    )

    doc.add_heading("9. Metrics and Reporting", level=1)
    doc.add_paragraph(
        "The SOC tracks the following metrics for all security incidents: Mean Time to Detect "
        "(MTTD), Mean Time to Respond (MTTR), Mean Time to Contain (MTTC), Mean Time to Recover, "
        "number of incidents by severity and type, percentage of incidents detected by automated "
        "monitoring versus human reporting, and false positive rate for automated alerts. These "
        "metrics are reported monthly to the CISO and quarterly to the executive team and board "
        "of directors."
    )

    doc.save(OUTPUT_DIR / "incident_response_procedure.docx")
    print("  Created: incident_response_procedure.docx")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Generating test corpus in {OUTPUT_DIR}/")
    generate_it_security_policy()
    generate_onboarding_guide()
    generate_expense_policy()
    generate_remote_work_policy()
    generate_incident_response()
    print(f"\nDone! Generated 5 documents in {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()