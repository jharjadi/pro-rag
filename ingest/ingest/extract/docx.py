"""DOCX extractor — produces structured blocks from .docx files.

Blocks: {type: heading|paragraph|table|list, text: str, meta: dict}
Tables are preserved as markdown (never shredded).
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.table import Table as DocxTable

from ingest.extract import Block

logger = logging.getLogger(__name__)


def _table_to_markdown(table: DocxTable) -> str:
    """Convert a python-docx Table to a markdown table string."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    lines: list[str] = []
    # Header row
    lines.append("| " + " | ".join(rows[0]) + " |")
    # Separator
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Data rows
    for row in rows[1:]:
        # Pad row to match header length if needed
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[: len(rows[0])]) + " |")

    return "\n".join(lines)


def _is_list_paragraph(paragraph) -> bool:
    """Check if a paragraph is a list item (has numbering/bullet or list style)."""
    # Check style name for list patterns
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name and any(
        kw in style_name.lower()
        for kw in ("list", "bullet", "number")
    ):
        return True

    # Check XML numPr element (explicit numbering)
    pPr = paragraph._element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
    )
    if pPr is None:
        return False
    numPr = pPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
    )
    return numPr is not None


def extract_docx(file_path: str | Path) -> list[Block]:
    """Extract structured blocks from a DOCX file.

    Args:
        file_path: Path to the .docx file.

    Returns:
        List of Block objects in document order.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if not file_path.suffix.lower() == ".docx":
        raise ValueError(f"Expected .docx file, got: {file_path.suffix}")

    doc = Document(str(file_path))
    blocks: list[Block] = []

    # python-docx exposes body elements in order (paragraphs + tables interleaved)
    # We iterate over the XML body children to preserve ordering.
    from docx.oxml.ns import qn

    body = doc.element.body
    # Build lookup maps for paragraphs and tables by their XML element
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    current_heading_level = 0

    for child in body:
        tag = child.tag

        if tag == qn("w:p"):
            para = para_map.get(child)
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # Detect headings
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    level = 1
                current_heading_level = level
                blocks.append(Block(
                    type="heading",
                    text=text,
                    meta={"level": level},
                ))
            elif _is_list_paragraph(para):
                blocks.append(Block(
                    type="list",
                    text=text,
                    meta={},
                ))
            else:
                blocks.append(Block(
                    type="paragraph",
                    text=text,
                    meta={},
                ))

        elif tag == qn("w:tbl"):
            table = table_map.get(child)
            if table is None:
                continue

            md = _table_to_markdown(table)
            if md.strip():
                num_rows = len(table.rows)
                num_cols = len(table.columns) if table.columns else 0
                blocks.append(Block(
                    type="table",
                    text=md,
                    meta={
                        "format": "markdown",
                        "num_rows": num_rows,
                        "num_cols": num_cols,
                    },
                ))

    logger.info(
        "Extracted %d blocks from %s (headings=%d, paragraphs=%d, tables=%d, lists=%d)",
        len(blocks),
        file_path.name,
        sum(1 for b in blocks if b.type == "heading"),
        sum(1 for b in blocks if b.type == "paragraph"),
        sum(1 for b in blocks if b.type == "table"),
        sum(1 for b in blocks if b.type == "list"),
    )

    return blocks
