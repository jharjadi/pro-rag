"""Tests for DOCX extractor.

Creates real .docx files in a temp directory to test extraction.
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from ingest.extract import Block
from ingest.extract.docx import extract_docx


@pytest.fixture
def tmp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


def _create_simple_docx(path: Path) -> Path:
    """Create a simple DOCX with headings and paragraphs."""
    doc = Document()
    doc.add_heading("Chapter 1", level=1)
    doc.add_paragraph("This is the first paragraph of chapter 1.")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_heading("Section 1.1", level=2)
    doc.add_paragraph("Content under section 1.1.")
    doc.add_heading("Chapter 2", level=1)
    doc.add_paragraph("Content of chapter 2.")
    filepath = path / "simple.docx"
    doc.save(str(filepath))
    return filepath


def _create_table_docx(path: Path) -> Path:
    """Create a DOCX with a table."""
    doc = Document()
    doc.add_heading("Data Report", level=1)
    doc.add_paragraph("Below is the data table:")
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    # Header
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Age"
    table.cell(0, 2).text = "City"
    # Data
    data = [
        ("Alice", "30", "New York"),
        ("Bob", "25", "Los Angeles"),
        ("Charlie", "35", "Chicago"),
    ]
    for i, (name, age, city) in enumerate(data, start=1):
        table.cell(i, 0).text = name
        table.cell(i, 1).text = age
        table.cell(i, 2).text = city
    doc.add_paragraph("End of report.")
    filepath = path / "with_table.docx"
    doc.save(str(filepath))
    return filepath


def _create_list_docx(path: Path) -> Path:
    """Create a DOCX with list items."""
    doc = Document()
    doc.add_heading("Shopping List", level=1)
    doc.add_paragraph("Apples", style="List Bullet")
    doc.add_paragraph("Bananas", style="List Bullet")
    doc.add_paragraph("Oranges", style="List Bullet")
    doc.add_paragraph("That's all.")
    filepath = path / "with_list.docx"
    doc.save(str(filepath))
    return filepath


# ── Tests ────────────────────────────────────────────────


class TestDocxExtractor:
    def test_simple_document(self, tmp_dir):
        filepath = _create_simple_docx(tmp_dir)
        blocks = extract_docx(filepath)

        assert len(blocks) > 0

        # Check headings
        headings = [b for b in blocks if b.type == "heading"]
        assert len(headings) == 3  # Chapter 1, Section 1.1, Chapter 2

        # Check heading levels
        assert headings[0].text == "Chapter 1"
        assert headings[0].meta["level"] == 1
        assert headings[1].text == "Section 1.1"
        assert headings[1].meta["level"] == 2
        assert headings[2].text == "Chapter 2"
        assert headings[2].meta["level"] == 1

        # Check paragraphs
        paragraphs = [b for b in blocks if b.type == "paragraph"]
        assert len(paragraphs) >= 3

    def test_table_extraction(self, tmp_dir):
        filepath = _create_table_docx(tmp_dir)
        blocks = extract_docx(filepath)

        # Should have a table block
        tables = [b for b in blocks if b.type == "table"]
        assert len(tables) == 1

        table = tables[0]
        assert "Name" in table.text
        assert "Alice" in table.text
        assert "Bob" in table.text
        assert table.meta["format"] == "markdown"
        assert table.meta["num_rows"] == 4
        assert table.meta["num_cols"] == 3

    def test_table_is_markdown_format(self, tmp_dir):
        filepath = _create_table_docx(tmp_dir)
        blocks = extract_docx(filepath)
        table = [b for b in blocks if b.type == "table"][0]

        lines = table.text.split("\n")
        # Should have header, separator, and data rows
        assert len(lines) >= 4  # header + separator + 3 data rows
        assert lines[0].startswith("|")
        assert "---" in lines[1]

    def test_list_extraction(self, tmp_dir):
        filepath = _create_list_docx(tmp_dir)
        blocks = extract_docx(filepath)

        lists = [b for b in blocks if b.type == "list"]
        assert len(lists) == 3
        assert lists[0].text == "Apples"
        assert lists[1].text == "Bananas"
        assert lists[2].text == "Oranges"

    def test_document_order_preserved(self, tmp_dir):
        filepath = _create_table_docx(tmp_dir)
        blocks = extract_docx(filepath)

        # Order should be: heading, paragraph, table, paragraph
        types = [b.type for b in blocks]
        assert types[0] == "heading"
        # Find table position
        table_idx = types.index("table")
        assert table_idx > 0  # Table comes after heading/paragraph

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            extract_docx("/nonexistent/file.docx")

    def test_wrong_extension(self, tmp_dir):
        filepath = tmp_dir / "test.txt"
        filepath.write_text("hello")
        with pytest.raises(ValueError, match="Expected .docx"):
            extract_docx(filepath)

    def test_empty_document(self, tmp_dir):
        doc = Document()
        filepath = tmp_dir / "empty.docx"
        doc.save(str(filepath))
        blocks = extract_docx(filepath)
        assert blocks == []
